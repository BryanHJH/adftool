import os
import sys
import subprocess
import zipfile
import logging
import re
import json

from flask import Flask, render_template, request, send_from_directory, send_file, current_app
from flask_sqlalchemy import SQLAlchemy
from werkzeug.utils import secure_filename
from werkzeug.datastructures import Headers
from io import BytesIO
from datetime import datetime
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from docx import Document
from html import escape

sys.path.insert(1, os.path.join(os.path.dirname(__file__), './modules'))
sys.path.insert(1, os.path.join(os.path.dirname(__file__), './bin'))
from magic_byte_analysis import analyze_file as magic_byte_analysis
from image_analyze import analyze_file as steganalysis
from pcap_analyze import analyze_pcap as pcapanalysis

# Application
app = Flask(__name__)
app.logger.setLevel(logging.INFO)

# Database
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///database.db'
db = SQLAlchemy(app)

# Directory to store all relevant files
UPLOAD_FOLDER = "/home/data"
RESULT_FOLDER = "/home/results"

class Scans(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(50), nullable=False)
    resultpath = db.Column(db.String(), nullable=False) # Storing the file path instead of the file itself
    fileextension = db.Column(db.String(50))
    date = db.Column(db.DateTime, default=datetime.utcnow)

    def __repr__(self):
        return '<Scan %r>' % self.id
    
def sanitize_for_xml(text):
    # Remove NULL bytes
    text = text.replace('\x00', '')
    # Replace other control characters with space
    text = re.sub(r'[\x01-\x1F\x7F]', ' ', text)
    return text

def clean_content_for_pdf(content):
    content = re.sub('<[^<]+?>', '', content)
    content = escape(content)
    content = re.sub(r'\n+', '\n', content)
    return content

def get_result_files(result_path):
    result_files = []
    file_contents = {}
    images = []

    for root, dirs, files in os.walk(result_path):
        for file in files:
            file_path = os.path.join(root, file)
            relative_path = os.path.relpath(file_path, result_path)
            result_files.append(relative_path)

            if file.endswith('.txt') or file.endswith('.bin'):
                with open(file_path, 'rb') as f:
                    content = f.read().decode('utf-8', 'ignore')
                    if content.strip():
                        file_contents[relative_path] = content

            elif file.endswith(('.jpg', '.jpeg', '.png', '.gif')):
                images.append(relative_path)

    return result_files, file_contents, images

#  Index page
@app.route("/", methods=["POST", "GET"])
def index():
    if request.method != "POST":
        return render_template("index.html", scans=Scans.query.order_by(Scans.date).all())
    
    file = request.files["file"]
    filename = secure_filename(file.filename)
    file_path = os.path.join(UPLOAD_FOLDER, filename) # The path to the file
    result_path = os.path.join(RESULT_FOLDER, f"results_{os.path.basename(file_path)}") # The path to the result

    progress_messages = ""

    try:
        file.save(file_path)
    except Exception as e:
        return f"Error: {e}"

    try:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        # Perform magic byte analysis and file signature checks using analyze_file
        is_match, signature_message, magic_bytes_message, file_extension, has_extension = magic_byte_analysis(file_path, verbose=False)

        if has_extension and is_match:
            if file_extension.lower() in ["png", "jpg", "jpeg", "gif", "bmp"]:
                os.makedirs(result_path, exist_ok=True)
                try:
                    # Performing steganalysis if the uploaded file is an image file and passes the magic bytes analysis check
                    subprocesses, progress_messages = steganalysis(file_path, verbose=False)
                except subprocess.CalledProcessError as e:
                    return f"Error executing steganalysis script: {str(e)}\nReturn code: {e.returncode}\nOutput: {e.output}"
                except Exception as e:
                    return f"Error during steganalysis: {str(e)}"
            elif file_extension.lower() in ["pcap", "pcapng"]:
                try:
                    # Performing pcap analysis using tshark if the uploaded file is a PCAP or PCAPNG file and passes the magic bytes analysis check
                    subprocesses, progress_messages = pcapanalysis(file_path, verbose=False)
                except subprocess.CalledProcessError as e:
                    return f"Error executing PCAP analysis script: {str(e)}\nReturn code: {e.returncode}\nOutput: {e.output}"
                except Exception as e:
                    return f"Error during PCAP analysis: {str(e)}"
                
        scan = Scans(filename=filename, resultpath=result_path, fileextension=file_extension)

        try:
            db.session.add(scan)
            db.session.commit()
            return render_template("index.html", scans=Scans.query.order_by(Scans.date).all(), signature_message=signature_message, magic_bytes_message=magic_bytes_message, progress_messages=progress_messages)
        except Exception as e:
            db.session.rollback()
            return f"Error adding scan to the database: {str(e)}. There was an issue adding your scan"
    except FileNotFoundError as e:
        return f"Error: str(e).\nThe uploaded file was not found."
    except Exception as e:
        return f"Error analyzing file: {str(e)}.\nThere was an issue analyzing the file.\n{list(filter(os.path.isfile, os.listdir('/home/data')))}"
    
# Viewing the record
@app.route("/view/<int:id>")
def view(id):
    scan = Scans.query.get_or_404(id)
    result_files, file_contents, images = get_result_files(scan.resultpath)
    result_files.sort()
    return render_template("view.html", scan=scan, result_files=result_files, file_contents=file_contents, images=images)

@app.route("/download_result/<int:scan_id>/<path:filename>")
def download_result(scan_id, filename):
    scan = Scans.query.get_or_404(scan_id)
    return send_from_directory(scan.resultpath, filename, as_attachment=True)

@app.route("/download_all/<int:scan_id>", methods=["POST"])
def download_all_results(scan_id):
    scan = Scans.query.get_or_404(scan_id)
    result_files, _, _ = get_result_files(scan.resultpath)

    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w') as zip_file:
        for file in result_files:
            file_path = os.path.join(scan.resultpath, file)
            zip_file.write(file_path, file)

    zip_buffer.seek(0)
    return send_file(zip_buffer, as_attachment=True, download_name=f"{scan.filename}_results.zip")

@app.route("/download_report/<int:scan_id>", methods=["POST"])
def download_report(scan_id):
    try:
        current_app.logger.info(f"Starting report generation for scan_id: {scan_id}")
        scan = Scans.query.get_or_404(scan_id)
        result_files, file_contents, images = get_result_files(scan.resultpath)
        report_format = request.form.get('report_format', 'pdf')
        filtered_content = request.form.get('filtered_content', None)

        if filtered_content:
            file_contents = json.loads(filtered_content)

        current_app.logger.info(f"Generating {report_format} report")

        if report_format == 'pdf':
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
            styles = getSampleStyleSheet()
            styles.add(ParagraphStyle(name='Small', fontSize=8, leading=10))
            story = []

            story.append(Paragraph(f"Analysis Results for {scan.filename}", styles['Title']))
            story.append(Paragraph(f"File Type: {scan.fileextension}", styles['Normal']))
            story.append(Paragraph(f"Analysis Date: {scan.date.strftime('%Y-%m-%d')}", styles['Normal']))
            story.append(Spacer(1, 12))

            for file, content in file_contents.items():
                story.append(Paragraph(f"File: {file}", styles['Heading2']))
                cleaned_content = clean_content_for_pdf(content)
                story.append(Paragraph(cleaned_content, styles['Small']))
                story.append(Spacer(1, 12))

            doc.build(story)
            buffer.seek(0)
            current_app.logger.info("PDF report generated successfully")
            
            headers = Headers()
            headers.add('Content-Disposition', 'attachment', filename=f"{scan.filename}_report.pdf")
            return current_app.response_class(buffer, mimetype='application/pdf', headers=headers)

        elif report_format == 'docx':
            buffer = BytesIO()
            doc = Document()
            doc.add_heading(f"Analysis Results for {scan.filename}", level=1)
            doc.add_paragraph(f"File Type: {scan.fileextension}")
            doc.add_paragraph(f"Analysis Date: {scan.date.strftime('%Y-%m-%d')}")

            for file, content in file_contents.items():
                doc.add_heading(f"File: {file}", level=2)
                doc.add_paragraph(sanitize_for_xml(content))

            doc.save(buffer)
            buffer.seek(0)
            current_app.logger.info("DOCX report generated successfully")
            
            headers = Headers()
            headers.add('Content-Disposition', 'attachment', filename=f"{scan.filename}_report.docx")
            return current_app.response_class(buffer, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', headers=headers)

        else:
            current_app.logger.warning(f"Invalid report format requested: {report_format}")
            return "Invalid report format", 400

    except Exception as e:
        current_app.logger.error(f"Error in download_report: {str(e)}", exc_info=True)
        return f"An error occurred: {str(e)}", 500
@app.route("/image/<int:scan_id>/<path:filename>")
def get_image(scan_id, filename):
    scan = Scans.query.get_or_404(scan_id)
    image_path = os.path.join(scan.resultpath, filename)
    return send_file(image_path)

if __name__ in "__main__":
    with app.app_context():
        db.drop_all()
        db.create_all()
    app.run(host='0.0.0.0', port=5001)